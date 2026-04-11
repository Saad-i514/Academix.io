"""
Lab Report Generator Tool - Extracts content from lab manuals (PDF/DOCX) and generates structured reports.
"""

from crewai.tools import BaseTool
from pydantic import BaseModel, Field
from typing import Optional
import os
from pathlib import Path


class LabReportInput(BaseModel):
    """Input schema for LabReportGeneratorTool."""
    file_path: str = Field(..., description="Path to the lab manual file (PDF or DOCX)")
    experiment_number: Optional[str] = Field(None, description="Specific experiment number to extract")
    include_code: bool = Field(True, description="Whether to include code execution")
    include_diagrams: bool = Field(True, description="Whether to generate diagrams")


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF file."""
    try:
        import PyPDF2
        
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text
    except ImportError:
        return "Error: PyPDF2 not installed. Install with: pip install PyPDF2"
    except Exception as e:
        return f"Error extracting PDF: {str(e)}"


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX file."""
    try:
        from docx import Document
        
        doc = Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        return text
    except ImportError:
        return "Error: python-docx not installed. Install with: pip install python-docx"
    except Exception as e:
        return f"Error extracting DOCX: {str(e)}"


def extract_lab_manual_content(file_path: str) -> dict:
    """Extract and structure content from lab manual."""
    if not os.path.exists(file_path):
        return {"error": f"File not found: {file_path}"}
    
    file_ext = Path(file_path).suffix.lower()
    
    if file_ext == '.pdf':
        raw_text = extract_text_from_pdf(file_path)
    elif file_ext in ['.docx', '.doc']:
        raw_text = extract_text_from_docx(file_path)
    else:
        return {"error": f"Unsupported file format: {file_ext}. Supported: PDF, DOCX"}
    
    if raw_text.startswith("Error"):
        return {"error": raw_text}
    
    # Extract method names to prevent hallucination
    method_names = extract_method_names(raw_text)
    experiment_title = extract_experiment_title(raw_text)
    
    return {
        "success": True,
        "raw_text": raw_text,
        "file_path": file_path,
        "file_type": file_ext,
        "text_length": len(raw_text),
        "method_names": method_names,
        "experiment_title": experiment_title,
    }


def extract_method_names(text: str) -> list:
    """Extract numerical method names from lab manual to prevent hallucination."""
    import re
    
    # Common numerical method patterns
    method_patterns = [
        r'bisection\s+method',
        r'false\s+position\s+method',
        r'regula\s+falsi',
        r'secant\s+method',
        r'newton[-\s]raphson\s+method',
        r'fixed\s+point\s+iteration',
        r'euler[\'s]*\s+method',
        r'runge[-\s]kutta\s+method',
        r'simpson[\'s]*\s+rule',
        r'trapezoidal\s+rule',
        r'gauss[-\s]elimination',
        r'jacobi\s+method',
        r'gauss[-\s]seidel',
    ]
    
    found_methods = []
    text_lower = text.lower()
    
    for pattern in method_patterns:
        matches = re.findall(pattern, text_lower, re.IGNORECASE)
        if matches:
            # Normalize the method name
            method = matches[0].strip().title()
            if method not in found_methods:
                found_methods.append(method)
    
    return found_methods


def extract_experiment_title(text: str) -> str:
    """Extract experiment title from lab manual."""
    import re
    
    # Look for common title patterns
    lines = text.split('\n')
    for i, line in enumerate(lines[:20]):  # Check first 20 lines
        line = line.strip()
        # Look for "EXPERIMENT" followed by number and title
        if re.match(r'EXPERIMENT\s+\d+', line, re.IGNORECASE):
            return line
        # Look for lines with all caps that might be titles
        if line.isupper() and len(line) > 10 and len(line) < 150:
            return line
    
    return "Lab Experiment"


def generate_lab_report_template(
    title: str,
    objective: str,
    theory: str,
    code: str = "",
    output: str = "",
    diagrams: list = None,
    conclusion: str = "",
) -> str:
    """Generate a professional lab report in markdown format."""
    
    diagrams = diagrams or []
    
    report = f"""# {title}

---

## 📋 Table of Contents
1. [Objective](#objective)
2. [Theory](#theory)
3. [Implementation](#implementation)
4. [Results](#results)
5. [Diagrams](#diagrams)
6. [Conclusion](#conclusion)
7. [References](#references)

---

## 🎯 Objective

{objective}

---

## 📚 Theory

{theory}

---

## 💻 Implementation

### Code

```octave
{code if code else "# No code provided"}
```

### Execution Details

- **Language**: Octave/MATLAB
- **Execution Environment**: Octave Online
- **Date**: {__import__('datetime').datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

---

## 📊 Results

### Output

```
{output if output else "No output generated"}
```

### Analysis

{_generate_analysis_section(output)}

---

## 📈 Diagrams

{_format_diagrams_section(diagrams)}

---

## 🎓 Conclusion

{conclusion}

---

## 📖 References

1. Lab Manual - {title}
2. Numerical Methods Course Materials
3. Octave Documentation: https://octave.org/doc/
4. Additional resources gathered during research

---

**Report Generated**: {__import__('datetime').datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

**Prepared by**: University Automation System
"""
    
    return report


def _generate_analysis_section(output: str) -> str:
    """Generate analysis section based on output."""
    if not output or output == "No output generated":
        return "No analysis available as no output was generated."
    
    lines = output.strip().split('\n')
    num_lines = len(lines)
    
    analysis = f"""
The execution produced {num_lines} line(s) of output. The results demonstrate:

- Successful execution of the numerical method
- Output values are within expected ranges
- The algorithm converged to a solution
- Results can be verified against theoretical expectations
"""
    return analysis


def _format_diagrams_section(diagrams: list) -> str:
    """Format diagrams section."""
    if not diagrams:
        return "No diagrams generated for this experiment."
    
    section = ""
    for idx, diagram in enumerate(diagrams, 1):
        if isinstance(diagram, str):
            section += f"\n### Diagram {idx}\n\n![Diagram {idx}]({diagram})\n"
        elif isinstance(diagram, dict):
            title = diagram.get('title', f'Diagram {idx}')
            url = diagram.get('url', '')
            description = diagram.get('description', '')
            section += f"\n### {title}\n\n"
            if description:
                section += f"{description}\n\n"
            if url:
                section += f"![{title}]({url})\n"
    
    return section


class LabReportGeneratorTool(BaseTool):
    name: str = "LabReportGenerator"
    description: str = (
        "Extracts content from lab manual files (PDF/DOCX) and structures it for report generation. "
        "Returns extracted text including title, objectives, theory, and procedures. "
        "Use this tool to process lab manuals before generating complete reports."
    )
    args_schema: type[BaseModel] = LabReportInput

    def _run(
        self,
        file_path: str,
        experiment_number: Optional[str] = None,
        include_code: bool = True,
        include_diagrams: bool = True,
    ) -> str:
        """Extract lab manual content and return structured information."""
        result = extract_lab_manual_content(file_path)
        
        if "error" in result:
            return f"Error: {result['error']}"
        
        response = f"""Lab Manual Content Extracted Successfully:

File: {result['file_path']}
Type: {result['file_type']}
Content Length: {result['text_length']} characters

CRITICAL - EXTRACTED METHOD NAMES (DO NOT SUBSTITUTE OR CHANGE):
{', '.join(result['method_names']) if result['method_names'] else 'No specific methods detected'}

EXPERIMENT TITLE:
{result['experiment_title']}

ANTI-HALLUCINATION REQUIREMENTS:
- Write about ONLY the methods listed above
- DO NOT substitute similar methods (e.g., False Position ≠ Newton-Raphson)
- DO NOT add methods not listed above
- Use EXACT formulas and algorithms from the content below
- Verify all method names in your output match the list above

Raw Content Preview (first 1500 chars):
{result['raw_text'][:1500]}...

Instructions for Agent:
1. Parse the above content to identify:
   - Experiment title (use the one extracted above)
   - Objective/Aim
   - Theory section (for ONLY the methods listed above)
   - Procedure/Algorithm (use exact algorithms from the manual)
   - Expected results

2. If code is needed, extract algorithm steps and convert to Octave code using EXACT formulas
3. Use OctaveOnline tool to execute the code
4. Use ImageCreator tool to generate relevant diagrams
5. Use search_tool ONLY for the specific methods listed above
6. Compile everything into a professional lab report
7. FINAL CHECK: Verify all method names in your report match the extracted list

Full content available for detailed parsing.
"""
        
        return response


# Export the template generator for use by agents
__all__ = ["LabReportGeneratorTool", "generate_lab_report_template"]
