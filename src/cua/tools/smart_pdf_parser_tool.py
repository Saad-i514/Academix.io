"""
SmartPDFParserTool — Extracts structured content from PDF/DOCX lab manuals.

Priority order:
  1. Adobe PDF Extract API  (ADOBE_CLIENT_ID + ADOBE_CLIENT_SECRET)
     — Best-in-class: extracts text, tables, figures, reading order, equations
     — Free: 500 documents/month
  2. pdfplumber             (no key, already installed)
     — Good for text-based PDFs, extracts tables
  3. PyPDF2                 (no key, fallback)
     — Basic text extraction only

Automatically identifies sections: Objective, Theory, Procedure, Results, etc.
"""

import os
import re
import json
import zipfile
import tempfile
import requests
from pathlib import Path
from pydantic import BaseModel, Field
from crewai.tools import BaseTool


class SmartPDFInput(BaseModel):
    file_path: str = Field(..., description="Absolute path to the PDF or DOCX file to parse")
    extract_tables: bool = Field(True, description="Whether to extract tables as structured Markdown")


# ── Section detection ─────────────────────────────────────────────────────────
SECTION_PATTERNS = {
    "objective":  re.compile(r"^(objective|aim|purpose|goal)\s*[:\-]?", re.I),
    "theory":     re.compile(r"^(theory|background|introduction|literature)\s*[:\-]?", re.I),
    "procedure":  re.compile(r"^(procedure|methodology|method|steps|algorithm)\s*[:\-]?", re.I),
    "equipment":  re.compile(r"^(equipment|tools|requirements|materials)\s*[:\-]?", re.I),
    "results":    re.compile(r"^(results|output|observations|findings)\s*[:\-]?", re.I),
    "conclusion": re.compile(r"^(conclusion|summary|discussion)\s*[:\-]?", re.I),
    "references": re.compile(r"^(references|bibliography|sources)\s*[:\-]?", re.I),
}


# ── Adobe PDF Extract ─────────────────────────────────────────────────────────

def _get_adobe_token() -> str | None:
    """Get Adobe IMS access token using client credentials."""
    client_id     = os.getenv("ADOBE_CLIENT_ID", "").strip()
    client_secret = os.getenv("ADOBE_CLIENT_SECRET", "").strip()
    if not client_id or not client_secret:
        return None
    try:
        resp = requests.post(
            "https://ims-na1.adobelogin.com/ims/token/v3",
            data={
                "grant_type":    "client_credentials",
                "client_id":     client_id,
                "client_secret": client_secret,
                "scope":         "openid,AdobeID,DCAPI",
            },
            timeout=15,
        )
        if resp.status_code == 200:
            return resp.json().get("access_token")
        return None
    except Exception:
        return None


def _extract_with_adobe(file_path: str) -> dict | None:
    """
    Use Adobe PDF Extract API to get structured JSON output.
    Returns structured sections dict or None if unavailable.
    """
    token = _get_adobe_token()
    if not token:
        return None

    client_id = os.getenv("ADOBE_CLIENT_ID", "").strip()
    headers = {
        "Authorization": f"Bearer {token}",
        "x-api-key":     client_id,
    }

    try:
        # Step 1: Create upload pre-signed URI
        resp = requests.post(
            "https://pdf-services.adobe.io/assets",
            headers={**headers, "Content-Type": "application/json"},
            json={"mediaType": "application/pdf"},
            timeout=15,
        )
        if resp.status_code != 200:
            return None

        asset_data   = resp.json()
        upload_uri   = asset_data.get("uploadUri")
        asset_id     = asset_data.get("assetID")

        # Step 2: Upload the PDF
        with open(file_path, "rb") as f:
            upload_resp = requests.put(
                upload_uri,
                data=f,
                headers={"Content-Type": "application/pdf"},
                timeout=60,
            )
        if upload_resp.status_code not in (200, 201):
            return None

        # Step 3: Create extraction job
        job_resp = requests.post(
            "https://pdf-services.adobe.io/operation/extractpdf",
            headers={**headers, "Content-Type": "application/json"},
            json={
                "assetID": asset_id,
                "getCharBounds": False,
                "includeStyling": False,
                "elementsToExtract": ["text", "tables"],
            },
            timeout=15,
        )
        if job_resp.status_code not in (200, 201):
            return None

        # Step 4: Poll for result
        job_location = job_resp.headers.get("Location") or job_resp.json().get("location")
        if not job_location:
            return None

        import time
        for _ in range(20):
            time.sleep(3)
            poll = requests.get(job_location, headers=headers, timeout=15)
            status = poll.json().get("status")
            if status == "done":
                result_url = poll.json().get("content", {}).get("downloadUri")
                break
            elif status == "failed":
                return None
        else:
            return None

        # Step 5: Download and parse result ZIP
        zip_resp = requests.get(result_url, timeout=30)
        if zip_resp.status_code != 200:
            return None

        with tempfile.TemporaryDirectory() as tmp:
            zip_path = os.path.join(tmp, "result.zip")
            with open(zip_path, "wb") as f:
                f.write(zip_resp.content)

            with zipfile.ZipFile(zip_path, "r") as z:
                z.extractall(tmp)

            json_path = os.path.join(tmp, "structuredData.json")
            if not os.path.exists(json_path):
                return None

            with open(json_path, "r", encoding="utf-8") as f:
                data = json.load(f)

        return _parse_adobe_json(data)

    except Exception:
        return None


def _parse_adobe_json(data: dict) -> dict:
    """Parse Adobe's structured JSON into our section format."""
    sections = {k: [] for k in SECTION_PATTERNS}
    sections["raw_text"] = []
    sections["tables"] = []
    current_section = "raw_text"

    elements = data.get("elements", [])
    for el in elements:
        el_type = el.get("Path", "")
        text    = el.get("Text", "").strip()

        if not text:
            # Handle tables
            if "Table" in el_type and el.get("filePaths"):
                # Table data is in a separate CSV — skip for now, use text fallback
                pass
            continue

        # Detect section from heading elements
        if "H1" in el_type or "H2" in el_type or "H3" in el_type:
            for sec_name, pattern in SECTION_PATTERNS.items():
                if pattern.match(text):
                    current_section = sec_name
                    break

        sections[current_section].append(text)
        sections["raw_text"].append(text)

    return {k: "\n".join(v) if isinstance(v, list) else v for k, v in sections.items()}


# ── pdfplumber fallback ───────────────────────────────────────────────────────

def _extract_with_pdfplumber(file_path: str, extract_tables: bool) -> dict:
    try:
        import pdfplumber
    except ImportError:
        return _extract_with_pypdf2(file_path)

    sections = {k: [] for k in SECTION_PATTERNS}
    sections["raw_text"] = []
    sections["tables"] = []
    current_section = "raw_text"

    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            if extract_tables:
                for tbl in page.extract_tables():
                    if tbl:
                        fmt = _format_table(tbl)
                        if fmt:
                            sections["tables"].append(fmt)

            text = page.extract_text(x_tolerance=3, y_tolerance=3) or ""
            for line in text.split("\n"):
                line = line.strip()
                if not line:
                    continue
                detected = False
                for sec_name, pattern in SECTION_PATTERNS.items():
                    if pattern.match(line):
                        current_section = sec_name
                        detected = True
                        break
                if not detected:
                    sections[current_section].append(line)
                sections["raw_text"].append(line)

    return {k: "\n".join(v) if isinstance(v, list) else v for k, v in sections.items()}


def _extract_with_pypdf2(file_path: str) -> dict:
    try:
        import PyPDF2
        text = ""
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                text += (page.extract_text() or "") + "\n"
        return {"raw_text": text, "tables": []}
    except Exception as e:
        return {"error": str(e)}


def _extract_docx(file_path: str, extract_tables: bool) -> dict:
    try:
        from docx import Document
    except ImportError:
        return {"error": "python-docx not installed"}

    doc = Document(file_path)
    sections = {k: [] for k in SECTION_PATTERNS}
    sections["raw_text"] = []
    sections["tables"] = []
    current_section = "raw_text"

    for para in doc.paragraphs:
        line = para.text.strip()
        if not line:
            continue
        detected = False
        for sec_name, pattern in SECTION_PATTERNS.items():
            if pattern.match(line):
                current_section = sec_name
                detected = True
                break
        if not detected:
            sections[current_section].append(line)
        sections["raw_text"].append(line)

    if extract_tables:
        for table in doc.tables:
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            fmt = _format_table(rows)
            if fmt:
                sections["tables"].append(fmt)

    return {k: "\n".join(v) if isinstance(v, list) else v for k, v in sections.items()}


def _format_table(rows: list) -> str:
    rows = [r for r in rows if any(str(c).strip() for c in r)]
    if not rows:
        return ""
    cols = max(len(r) for r in rows)
    rows = [list(r) + [""] * (cols - len(r)) for r in rows]
    header    = "| " + " | ".join(str(c) for c in rows[0]) + " |"
    separator = "| " + " | ".join(["---"] * cols) + " |"
    body      = "\n".join("| " + " | ".join(str(c) for c in row) + " |" for row in rows[1:])
    return f"{header}\n{separator}\n{body}"


# ── Tool ─────────────────────────────────────────────────────────────────────

class SmartPDFParserTool(BaseTool):
    name: str = "SmartPDFParser"
    description: str = (
        "Intelligently parses PDF and DOCX lab manuals using Adobe PDF Extract API (primary) "
        "or pdfplumber (fallback). Automatically identifies sections (Objective, Theory, "
        "Procedure, Results, etc.) and extracts tables as structured Markdown. "
        "Use this for any uploaded lab manual or assignment document."
    )
    args_schema: type[BaseModel] = SmartPDFInput

    def _run(self, file_path: str, extract_tables: bool = True) -> str:
        if not os.path.exists(file_path):
            return f"Error: File not found at '{file_path}'"

        ext = Path(file_path).suffix.lower()
        parser_used = "unknown"

        if ext in (".docx", ".doc"):
            data = _extract_docx(file_path, extract_tables)
            parser_used = "python-docx"
        elif ext == ".pdf":
            # Try Adobe first
            adobe_result = _extract_with_adobe(file_path)
            if adobe_result:
                data = adobe_result
                parser_used = "Adobe PDF Extract API"
            else:
                data = _extract_with_pdfplumber(file_path, extract_tables)
                parser_used = "pdfplumber"
        else:
            return f"Error: Unsupported format '{ext}'. Supported: .pdf, .docx"

        if "error" in data:
            return f"Extraction error: {data['error']}"

        # Build structured output
        output_parts = [f"File: {Path(file_path).name} | Parser: {parser_used}\n"]

        section_labels = {
            "objective":  "OBJECTIVE",
            "theory":     "THEORY / BACKGROUND",
            "procedure":  "PROCEDURE / METHODOLOGY",
            "equipment":  "EQUIPMENT / TOOLS",
            "results":    "RESULTS",
            "conclusion": "CONCLUSION",
            "references": "REFERENCES",
        }

        for key, label in section_labels.items():
            content = data.get(key, "").strip()
            if content:
                output_parts.append(f"=== {label} ===\n{content}\n")

        if data.get("tables"):
            output_parts.append("=== EXTRACTED TABLES ===")
            for i, tbl in enumerate(data["tables"] if isinstance(data["tables"], list) else [data["tables"]], 1):
                if tbl:
                    output_parts.append(f"\nTable {i}:\n{tbl}\n")

        # Fallback to raw text if no sections detected
        has_sections = any(data.get(k, "").strip() for k in section_labels)
        if not has_sections:
            raw = data.get("raw_text", "").strip()
            output_parts.append(f"=== FULL TEXT (no sections auto-detected) ===\n{raw[:4000]}")

        return "\n".join(output_parts)
