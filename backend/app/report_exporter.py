"""Professional DOCX exporter for Academix.io reports."""
import re
import io
from datetime import datetime


# ── Helpers ───────────────────────────────────────────────────────────────────

def _strip_md_links(text):
    return re.sub(r"\[([^\]]+)\]\([^)]*\)", r"\1", text)

def _latex_to_text(latex):
    t = re.sub(r"\$\$|\$|\\\(|\\\)|\\\[|\\\]", "", latex)
    t = re.sub(r"\\frac\{([^}]+)\}\{([^}]+)\}", r"(\1)/(\2)", t)
    t = re.sub(r"\\text\{([^}]+)\}", r"\1", t)
    t = re.sub(r"\\approx", "≈", t)
    t = re.sub(r"\\times", "×", t)
    t = re.sub(r"\\[a-zA-Z]+\{([^}]*)\}", r"\1", t)
    t = re.sub(r"\\[a-zA-Z]+", "", t)
    t = re.sub(r"[{}]", "", t)
    return t.strip()

def _strip_latex(text):
    text = re.sub(r"\$\$[\s\S]+?\$\$", lambda m: _latex_to_text(m.group(0)), text)
    text = re.sub(r"\$[^$\n]+?\$", lambda m: _latex_to_text(m.group(0)), text)
    text = re.sub(r"\\\([\s\S]+?\\\)", lambda m: _latex_to_text(m.group(0)), text)
    text = re.sub(r"\\\[[\s\S]+?\\\]", lambda m: _latex_to_text(m.group(0)), text)
    text = re.sub(r"\\frac\{([^}]+)\}\{([^}]+)\}", r"(\1)/(\2)", text)
    text = re.sub(r"\\text\{([^}]+)\}", r"\1", text)
    text = re.sub(r"\\approx", "≈", text)
    text = re.sub(r"\\times", "×", text)
    text = re.sub(r"\\[a-zA-Z]+\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\[a-zA-Z]+", "", text)
    return text

def _clean(text):
    return _strip_latex(_strip_md_links(text))

def _fix_code(code):
    return code.replace("\r\n", "\n").replace("\r", "\n")

def _heading_level(line):
    m = re.match(r"^(#{1,6})\s+(.*)", line)
    if m:
        return len(m.group(1)), m.group(2).strip()
    return 0, line

def _extract_title(md, fallback):
    skip = {"university", "course", "title page", "title", "report"}
    for line in md.split("\n"):
        if line.startswith("# "):
            c = line[2:].strip()
            if c.lower() not in skip and len(c) > 3:
                return c
        if "experiment title:" in line.lower():
            parts = line.split(":", 1)
            if len(parts) > 1:
                t = re.sub(r"^#+\s*", "", parts[1]).strip()
                if t:
                    return t
    return fallback


# ── DOCX Export ───────────────────────────────────────────────────────────────

def markdown_to_docx(markdown_text: str, title: str = "Academic Report") -> bytes:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    real_title = _extract_title(markdown_text, title)
    doc = Document()

    for sec in doc.sections:
        sec.top_margin    = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin   = Inches(1.25)
        sec.right_margin  = Inches(1.25)

    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(12)

    _title_page(doc, real_title)
    doc.add_page_break()

    lines = markdown_text.split("\n")
    i = 0
    while i < len(lines):
        raw = lines[i]
        level, htext = _heading_level(raw)

        if level == 1:
            p = doc.add_heading(_clean(htext), level=1)
            _style_h(p, 16, "1a1a2e")
        elif level == 2:
            p = doc.add_heading(_clean(htext), level=2)
            _style_h(p, 14, "1a1a2e")
        elif level == 3:
            p = doc.add_heading(_clean(htext), level=3)
            _style_h(p, 12, "4c1d95")
        elif level >= 4:
            p = doc.add_paragraph()
            r = p.add_run(_clean(htext))
            r.bold = True
            r.font.size = Pt(11)
            r.font.name = "Calibri"
        elif raw.strip() == "---":
            _hr(doc)
        elif raw.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            _code_block(doc, _fix_code("\n".join(code_lines)))
        elif raw.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[-| :]+\|", lines[i + 1]):
            tbl = []
            while i < len(lines) and lines[i].startswith("|"):
                tbl.append(lines[i])
                i += 1
            _table(doc, tbl)
            continue
        elif raw.startswith("- ") or raw.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            _inline(p, _clean(raw[2:].strip()))
        elif re.match(r"^\d+\.\s", raw):
            p = doc.add_paragraph(style="List Number")
            _inline(p, _clean(re.sub(r"^\d+\.\s", "", raw).strip()))
        elif not raw.strip():
            doc.add_paragraph("")
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = Pt(18)
            _inline(p, _clean(raw))

        i += 1

    _footer(doc)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _title_page(doc, title):
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    for _ in range(5):
        doc.add_paragraph("")

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(title)
    r.bold = True
    r.font.size = Pt(24)
    r.font.name = "Calibri"
    r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    _hr(doc, color="4c1d95", sz="10")

    doc.add_paragraph("")
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("Generated by Academix.io")
    rs.font.size = Pt(10)
    rs.italic = True
    rs.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    for _ in range(4):
        doc.add_paragraph("")

    for label, value in [
        ("Student Name:", "[Student Name]"),
        ("Date:", datetime.now().strftime("%B %d, %Y")),
        ("Instructor:", "[Instructor Name]"),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(f"{label}  ")
        r1.bold = True
        r1.font.size = Pt(12)
        r1.font.name = "Calibri"
        r2 = p.add_run(value)
        r2.font.size = Pt(12)
        r2.font.name = "Calibri"


def _style_h(paragraph, size, color_hex):
    from docx.shared import Pt, RGBColor
    rgb = tuple(int(color_hex[i:i+2], 16) for i in (0, 2, 4))
    for run in paragraph.runs:
        run.font.size = Pt(size)
        run.bold = True
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(*rgb)


def _inline(paragraph, text):
    from docx.shared import Pt
    pattern = re.compile(r"(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)")
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            paragraph.add_run(text[last:m.start()])
        if m.group(1).startswith("**"):
            r = paragraph.add_run(m.group(2))
            r.bold = True
        elif m.group(1).startswith("*"):
            r = paragraph.add_run(m.group(3))
            r.italic = True
        else:
            r = paragraph.add_run(m.group(4))
            r.font.name = "Courier New"
            r.font.size = Pt(10)
        last = m.end()
    if last < len(text):
        paragraph.add_run(text[last:])


def _code_block(doc, code):
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    for line in code.split("\n"):
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F0F0F5")
        pPr.append(shd)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line)
        r.font.name = "Courier New"
        r.font.size = Pt(9)


def _table(doc, table_lines):
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    rows = []
    for line in table_lines:
        if re.match(r"^\|[-| :]+\|$", line.strip()):
            continue
        cells = [_clean(c.strip()) for c in line.strip().strip("|").split("|")]
        rows.append(cells)

    if not rows:
        return

    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"

    for r_idx, row_data in enumerate(rows):
        for c_idx in range(cols):
            cell_text = row_data[c_idx] if c_idx < len(row_data) else ""
            cell = table.cell(r_idx, c_idx)
            if r_idx == 0:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "1a1a2e")
                tcPr.append(shd)
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            run.font.size = Pt(10)
            run.font.name = "Calibri"
            if r_idx == 0:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    doc.add_paragraph("")


def _hr(doc, color="AAAAAA", sz="6"):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _footer(doc):
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"Academix.io  •  {datetime.now().strftime('%B %d, %Y')}")
        r.font.size = Pt(9)
        r.font.name = "Calibri"
        r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
